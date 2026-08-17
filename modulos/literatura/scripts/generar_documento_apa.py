#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Compilador académico APA 7: convierte un borrador Markdown en .docx y/o .pdf.

Uso:
    generar_documento_apa.py borrador.md [--pdf] [--estimar] [--salida DIR]
    generar_documento_apa.py --plantilla
"""

import argparse
import os
import re
import shutil
import subprocess
import sys
from datetime import date

import mistune
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT = "Times New Roman"
SIZE = 12
MARGIN = Cm(2.54)
INDENT = Cm(1.27)
LINE = 2.0
CENTER = WD_ALIGN_PARAGRAPH.CENTER
LEFT = WD_ALIGN_PARAGRAPH.LEFT
RIGHT = WD_ALIGN_PARAGRAPH.RIGHT

BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

MESES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]

FRONT_RE = re.compile(r"^---\s*\n(.*?)\n---\s*\n?", re.S)

META_DEFAULT = {
    "titulo": "Trabajo académico",
    "autor": "Arturo Andrés Baena Arias",
    "universidad": "Universidad de Cartagena",
    "programa": "Programa de Administración de Empresas",
    "curso": "",
    "docente": "",
    "fecha": "",
}


def fecha_hoy():
    d = date.today()
    return f"{d.day} de {MESES[d.month - 1]} de {d.year}"


def parse_front_matter(text):
    meta = dict(META_DEFAULT)
    if not meta["fecha"]:
        meta["fecha"] = fecha_hoy()
    body = text
    m = FRONT_RE.match(text)
    if m:
        body = text[m.end():].lstrip()
        for line in m.group(1).splitlines():
            fm = re.match(r"^\s*([A-Za-zÁÉÍÓÚÑáéíóúñ]+)\s*:\s*(.+?)\s*$", line)
            if fm:
                meta[fm.group(1).lower()] = fm.group(2)
    if not meta["fecha"]:
        meta["fecha"] = fecha_hoy()
    return meta, body


def _set_run_font(run, bold=False, italic=False, size=SIZE):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)


def add_run_text(p, text, bold=False, italic=False):
    if not text:
        return
    run = p.add_run(text)
    _set_run_font(run, bold=bold, italic=italic)


def _set_outline(p, level):
    pPr = p._p.get_or_add_pPr()
    el = OxmlElement("w:outlineLvl")
    el.set(qn("w:val"), str(level))
    pPr.append(el)


def para(doc, align=None, left=None, first=None, before=0, after=0, outline=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if left is not None:
        pf.left_indent = left
    if first is not None:
        pf.first_line_indent = first
    if outline is not None:
        _set_outline(p, outline)
    return p


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = MARGIN
    sec.bottom_margin = MARGIN
    sec.left_margin = MARGIN
    sec.right_margin = MARGIN
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SIZE)
    try:
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    except AttributeError:
        pass
    pf = st.paragraph_format
    pf.line_spacing = LINE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return doc


def add_page_number(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = RIGHT
    p.paragraph_format.line_spacing = 1.0
    for run in p.runs:
        run.text = ""
    run = p.add_run()
    _set_run_font(run)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def add_cover(doc, meta):
    p = para(doc, align=CENTER, before=0)
    add_run_text(p, meta["titulo"], bold=True)
    para(doc)
    for key in ("autor", "universidad", "programa", "curso", "docente", "fecha"):
        value = meta.get(key, "").strip()
        if value:
            p = para(doc, align=CENTER)
            add_run_text(p, value)


def _walk_inline_into(p, node, bold=False, italic=False):
    for child in node.get("children", []):
        t = child["type"]
        raw = child.get("raw") or child.get("text") or ""
        if t == "text":
            add_run_text(p, raw, bold=bold, italic=italic)
        elif t == "strong":
            _walk_inline_into(p, child, bold=True, italic=italic)
        elif t == "emphasis":
            _walk_inline_into(p, child, bold=bold, italic=True)
        elif t == "codespan":
            add_run_text(p, raw, bold=bold, italic=italic)
        elif t in ("link", "image"):
            _walk_inline_into(p, child, bold=bold, italic=italic)
        else:
            _walk_inline_into(p, child, bold=bold, italic=italic)


def render_inline(p, nodes):
    for node in nodes:
        t = node["type"]
        raw = node.get("raw") or node.get("text") or ""
        if t == "text":
            add_run_text(p, raw)
        elif t == "strong":
            _walk_inline_into(p, node, bold=True)
        elif t == "emphasis":
            _walk_inline_into(p, node, italic=True)
        elif t == "codespan":
            add_run_text(p, raw)
        elif t == "link":
            _walk_inline_into(p, node)
        elif t in ("linebreak", "softbreak"):
            add_run_text(p, " ")
        else:
            _walk_inline_into(p, node)


def add_heading(doc, tok, level):
    style_map = {
        1: (CENTER, True, False),
        2: (LEFT, True, False),
        3: (LEFT, True, True),
    }
    align, bold, italic = style_map.get(level, (LEFT, True, False))
    p = para(doc, align=align, outline=level - 1)
    _walk_inline_into(p, {"children": tok.get("children", [])}, bold=bold, italic=italic)


def plain_text(tok):
    parts = []

    def rec(nodes):
        for node in nodes:
            raw = node.get("raw") or node.get("text") or ""
            if raw:
                parts.append(raw)
            rec(node.get("children", []))

    rec(tok.get("children", []))
    return " ".join(parts)


def add_list(doc, tok):
    ordered = tok.get("attrs", {}).get("ordered", False)
    num = 0
    for item in tok.get("children", []):
        num += 1
        p = para(doc, left=INDENT, first=Cm(-0.63))
        marker = f"{num}. " if ordered else "\u2022  "
        add_run_text(p, marker)
        for sub in item.get("children", []):
            if sub["type"] == "paragraph":
                render_inline(p, sub.get("children", []))
            elif sub["type"] in ("list", "blockquote"):
                render_inline(p, sub.get("children", []))


def add_body(doc, meta, tokens):
    doc.add_page_break()
    p = para(doc, align=CENTER)
    add_run_text(p, meta["titulo"], bold=True)
    para(doc)
    in_refs = False
    for tok in tokens:
        t = tok["type"]
        if t == "heading":
            level = tok.get("attrs", {}).get("level", 1)
            text = plain_text(tok).strip()
            if text.lower() == "referencias":
                doc.add_page_break()
                p = para(doc, align=CENTER)
                add_run_text(p, "Referencias", bold=True)
                in_refs = True
                continue
            add_heading(doc, tok, level)
        elif t == "paragraph":
            if in_refs:
                p = para(doc, left=INDENT, first=Cm(-1.27))
            else:
                p = para(doc, first=INDENT)
            render_inline(p, tok.get("children", []))
        elif t == "block_quote":
            p = para(doc, left=INDENT, first=Pt(0))
            for sub in tok.get("children", []):
                if sub["type"] == "paragraph":
                    render_inline(p, sub.get("children", []))
        elif t == "list":
            add_list(doc, tok)
        else:
            continue


def estimar(text):
    _, body = parse_front_matter(text)
    body = re.sub(r"```.*?```", " ", body, flags=re.S)
    body = re.sub(r"`[^`]*`", " ", body)
    body = re.sub(r"[#>*_\-\[\]()!]", " ", body)
    words = [w for w in body.split() if re.search(r"[A-Za-zÁÉÍÓÚÑáéíóúñ0-9]", w)]
    return len(words), len(words) / 250.0


def build_plantilla():
    doc = setup_document()
    add_page_number(doc.sections[0])
    meta = dict(META_DEFAULT)
    meta["fecha"] = fecha_hoy()
    add_cover(doc, meta)
    doc.add_page_break()
    p = para(doc, align=CENTER)
    add_run_text(p, "Título del trabajo", bold=True)
    para(doc)
    p = para(doc, first=INDENT)
    add_run_text(
        p,
        "Párrafo de ejemplo con cita parentética (Apellido, año). "
        "Este documento muestra la base de estilos APA 7 aplicada automáticamente.",
    )
    p = para(doc, left=INDENT, first=Cm(-1.27))
    add_run_text(p, "Apellido, A. (Año). Título de la obra. Editorial.")
    return doc


def main():
    ap = argparse.ArgumentParser(
        description="Compila un borrador Markdown a un documento APA 7 (.docx / .pdf)."
    )
    ap.add_argument("borrador", nargs="?", help="Ruta al archivo borrador.md (o usa --input)")
    ap.add_argument("--input", default=None, help="Ruta al archivo borrador.md (alternativa al posicional)")
    ap.add_argument("--output", default=None, help="Ruta del .docx de salida (por defecto, junto al borrador)")
    ap.add_argument("--pdf", action="store_true", help="Genera también el .pdf vía LibreOffice")
    ap.add_argument("--estimar", action="store_true", help="Cuenta palabras y estima páginas")
    ap.add_argument("--salida", default=None, help="Carpeta de salida (por defecto, junto al borrador)")
    ap.add_argument("--plantilla", action="store_true", help="Regenera plantillas/apa_base.docx")
    args = ap.parse_args()

    if args.plantilla:
        path = os.path.join(BASE, "plantillas", "apa_base.docx")
        build_plantilla().save(path)
        print(f"Plantilla generada: {path}")
        return

    borrador = args.input or args.borrador
    if not borrador:
        ap.error("Debes indicar el archivo borrador.md (posicional o --input) o usar --plantilla")

    with open(borrador, encoding="utf-8") as f:
        text = f.read()

    if args.estimar:
        n, pag = estimar(text)
        print(f"Palabras: {n} | Estimación: ~{pag:.1f} páginas (250 palabras/página)")
        return

    meta, body = parse_front_matter(text)
    md = mistune.create_markdown(renderer=None)
    tokens = md(body)

    doc = setup_document()
    add_page_number(doc.sections[0])
    add_cover(doc, meta)
    add_body(doc, meta, tokens)

    if args.output:
        out_docx = args.output
        if not out_docx.lower().endswith(".docx"):
            out_docx += ".docx"
        src_dir = os.path.dirname(os.path.abspath(out_docx))
        base = os.path.splitext(os.path.basename(out_docx))[0]
    else:
        src_dir = args.salida or os.path.dirname(os.path.abspath(borrador))
        base = os.path.splitext(os.path.basename(borrador))[0]
        out_docx = os.path.join(src_dir, base + ".docx")
    os.makedirs(src_dir, exist_ok=True)
    doc.save(out_docx)
    print(f"DOCX generado: {out_docx}")

    if args.pdf:
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            sys.exit("LibreOffice no encontrado; no se pudo generar el PDF.")
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", src_dir, out_docx],
            check=True,
        )
        print(f"PDF generado: {os.path.join(src_dir, base + '.pdf')}")


if __name__ == "__main__":
    main()